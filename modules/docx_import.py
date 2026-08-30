"""
modules/docx_import.py
Import a Select AI Agent Builder project from a Word (.docx) configuration document.

Reads the standard Select AI Agent configuration document format:
  - Two-column tables (Field | Value) for all project settings
  - Multi-line text fields (Agent Instructions, Task Instruction, etc.)
  - SQL COMMENT blocks (NL2SQL comments)

The document is parsed into the same data dict format that project_import.parse_csv()
produces, then routed through the same build_project() / validate_project() pipeline.
This means both import paths produce identical project dicts and go through the
same Step 7 review before any SQL is generated.

Requires: python-docx  (pip install python-docx)
Called from: modules/conversation.py run_new() — option 3
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Optional

from core import state as state_module
from core import config as cfg_module
from modules.project_import import (
    _KEY_ALIASES, _norm_key, build_project, validate_project,
    print_import_summary,
)


# ─────────────────────────────────────────────────────────────────────────────
# Additional key aliases specific to the Word doc format
# (supplements the shared _KEY_ALIASES in project_import.py)
# ─────────────────────────────────────────────────────────────────────────────

_DOCX_EXTRA_ALIASES: dict[str, str] = {
    # Word doc uses "SQL Profile Name", "SQL Tool Name" etc
    "sqlprofilename":           "nl2sql_profile",
    "nl2sqlprofilename":        "nl2sql_profile",
    "sqltoolname":              "sql_tool",
    "ragtoolname":              "rag_tool",
    "ragprofilename":           "rag_profile",
    "vectorindexname":          "vector_index",
    "agentname":                "agent_name",
    "taskname":                 "task_name",
    "teamname":                 "team_name",
    "dsnalias":                 "dsn",
    "ragsubjectmatter":         "rag_subject",
    "ragfiletypes":             "rag_file_types",
    "ragbucketname":            "rag_bucket",
    "ragobjectstorageurl":      "rag_url",
    "ragnamespace":             "rag_namespace",
    "chunksize":                "chunk_size",
    "chunkoverlap":             "chunk_overlap",
    "llmchatmodel":             "chat_model",
    "llmembedmodel":            "embed_model",
    "llmtemperature":           "llm_temperature",
    "llmmaxtokens":             "llm_max_tokens",
    "maxtokens":                "llm_max_tokens",
    "temperature":              "llm_temperature",
    "sqltargetschema":          "sql_schema",
    "sqltool":                  "sql_tool",
    "sqltoolprofile":           "nl2sql_profile",
    "underlyingtable1":         "sql_table_1",
    "underlyingtable2":         "sql_table_2",
    "nl2sqlobjectlist":         "sql_tables",
    "nl2sqltables":             "sql_tables",
    "nl2sqlunderlyingtables":   "sql_tables",
    "objectlist":               "sql_tables",
    # Shared error log table for builder-generated OML4Py custom tools —
    # explicit override so it can match an already-existing table (e.g. one
    # created by hand-written code pasted in via PL/SQL Body) instead of the
    # auto-derived "<SCHEMA>_ERROR_LOG" default. One per project, not per tool.
    "errorlogtablename":        "error_log_table",
    "errorlogtable":            "error_log_table",
    "agentdescription":         "agent_description",
    "agentinstructions":        "agent_instructions",
    "taskinstruction":          "task_instruction",
    "sqltooldescrip":           "sql_tool_description",
    "sqltooledescription":      "sql_tool_description",
    "sqltoodescription":        "sql_tool_description",
    "sqltooldescrip":           "sql_tool_description",
    "sqltooldescription":       "sql_tool_description",
    "ragtooldescrip":           "rag_tool_description",
    "ragtooldescription":       "rag_tool_description",
    "ragsubjectmatter":         "rag_subject",
    # NL2SQL comment blocks — the Word doc splits these into separate cells
    "nl2sqlcomments":                       "nl2sql_comments_sql",
    "nl2sqlcomments(tablelevel)":           "nl2sql_comments_sql",
    "nl2sqlcomments(columnlevel)":          "nl2sql_comments_sql_col",
    "nl2sqlmetadatacomments":               "nl2sql_comments_sql",
    "nl2sqlmetadatacomments(tablelevel)":   "nl2sql_comments_sql",
    "nl2sqlmetadatacomments(columnlevel)":  "nl2sql_comments_sql_col",
}

_ALL_ALIASES = {**_KEY_ALIASES, **_DOCX_EXTRA_ALIASES}


_NUMBERED_TABLE_PAT = re.compile(r"^(?:underlyingtable|sourcetable|nl2sqltable)(\d+)$")

# Canonical keys that hold literal code (PL/SQL, Python, SQL fragments) rather
# than prose. Word's AutoCorrect silently rewrites plain ASCII punctuation
# typed or pasted into these fields into "smart" Unicode look-alikes — an
# em dash for "--", curly quotes for straight quotes, a non-breaking space
# for a regular space. Invisible on screen, but when the resulting text is
# submitted as a SQL string literal it can corrupt the literal boundary
# (a stray em dash byte sequence has caused ORA-00984 "column not allowed
# here" — the parser sees the literal end early and reads trailing words as
# bare identifiers) or otherwise not mean what it looks like it means. Code
# fields get normalised back to plain ASCII; prose fields (Role, Instruction)
# are left untouched since a real em dash there is just typography.
_CODE_FIELD_PAT = re.compile(
    r"^custom_tool_\d+_(plsqlbody|code|pythonscript|dataquery|paramspec)$"
)
_SMART_PUNCT_MAP = {
    "\u2014": "--",   # em dash
    "\u2013": "-",    # en dash
    "\u2018": "'",    # left single quote
    "\u2019": "'",    # right single quote / apostrophe
    "\u201c": '"',    # left double quote
    "\u201d": '"',    # right double quote
    "\u00a0": " ",    # non-breaking space
}


def _normalize_code_punctuation(text: str) -> str:
    """Replace Word's smart-typography substitutions with plain ASCII."""
    for smart, plain in _SMART_PUNCT_MAP.items():
        text = text.replace(smart, plain)
    return text


# Custom Tool N <field> — e.g. "Custom Tool 1 Name", "Custom Tool 3 PL/SQL Body".
# Field suffix is normalised the same way as everything else (no spaces/punct),
# so "PL/SQL Body" -> "plsqlbody", "Function Name" -> "functionname", etc.
_NUMBERED_CUSTOM_TOOL_PAT = re.compile(
    r"^customtool(\d+)(name|functionname|instruction|plsqlbody|code|type|inputs"
    r"|pythonscript|dataquery|pyqscriptname|credentialname|parameters|paramspec)$"
)


def _resolve_key(raw: str) -> Optional[str]:
    """Resolve a raw field label to a canonical key using both alias tables.

    Falls back to two dynamic pattern matches when no direct alias exists:

    1. "Underlying Table N" (or "Source Table N" / "NL2SQL Table N") for ANY
       N — not just 1/2 — so a schema with 6+ source tables (e.g. ACME_CORP)
       works without a new hardcoded alias per table number.

    2. "Custom Tool N <field>" for ANY N and any of: Name, Function Name,
       Instruction, PL/SQL Body (or Code), Type, Inputs — letting a Word doc
       define an arbitrary number of custom function-type tools (PL/SQL or
       OML4Py-wrapped), each carrying its own already-tested implementation
       body rather than relying on a hardcoded template. See
       _extract_custom_tools() for how these are assembled.
    """
    norm = _norm_key(raw)
    direct = _ALL_ALIASES.get(norm)
    if direct:
        return direct
    m = _NUMBERED_TABLE_PAT.match(norm)
    if m:
        return f"sql_table_{m.group(1)}"
    m = _NUMBERED_CUSTOM_TOOL_PAT.match(norm)
    if m:
        return f"custom_tool_{m.group(1)}_{m.group(2)}"
    return None


# ─────────────────────────────────────────────────────────────────────────────
# Multi-line field detection
# ─────────────────────────────────────────────────────────────────────────────

# These canonical keys are expected to have long multi-line values.
# When we encounter them in a table, we preserve newlines.
_MULTILINE_KEYS = {
    "agent_instructions",
    "task_instruction",
    "sql_tool_description",
    "rag_tool_description",
    "nl2sql_comments",
    "nl2sql_comments_sql",
    "nl2sql_comments_sql_col",
}

# Fields that contain SQL COMMENT blocks — extracted separately
_SQL_COMMENT_FIELDS = {
    "nl2sql_comments",
    "nl2sql_comments_sql",
    "nl2sql_comments_sql_col",
    "nl2sql_clearing_comments",
}


# ─────────────────────────────────────────────────────────────────────────────
# Core docx extraction
# ─────────────────────────────────────────────────────────────────────────────

def _cell_text(cell) -> str:
    """Extract clean text from a docx table cell, preserving internal newlines.

    python-docx's para.text only returns text up to the first inline line
    break (<w:br/>) — everything after it is silently dropped. This matters
    for multi-line values like PL/SQL bodies and Python scripts that are stored
    as runs within a single paragraph (using add_run().add_break() rather than
    separate paragraphs). We walk the XML runs directly to reconstruct the
    full content, converting <w:br/> elements back to newlines.
    """
    from docx.oxml.ns import qn as _qn
    lines = []
    for para in cell.paragraphs:
        para_parts = []
        for child in para._p:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag == "r":  # w:r run element
                for sub in child:
                    sub_tag = sub.tag.split("}")[-1] if "}" in sub.tag else sub.tag
                    if sub_tag == "t":  # w:t text
                        para_parts.append(sub.text or "")
                    elif sub_tag == "br":  # w:br line break
                        para_parts.append("\n")
            elif tag == "hyperlink":
                # Hyperlinks also contain runs
                for run in child:
                    for sub in run:
                        sub_tag = sub.tag.split("}")[-1] if "}" in sub.tag else sub.tag
                        if sub_tag == "t":
                            para_parts.append(sub.text or "")
        para_text = "".join(para_parts)
        if para_text.strip():
            lines.append(para_text)
    return "\n".join(lines)


def _is_header_row(cells: list) -> bool:
    """Return True if this looks like a column header row (Field | Value | Status etc.)"""
    if not cells:
        return False
    first = cells[0].strip().lower()
    return first in ("field", "key", "name", "setting", "parameter", "column", "table")


def _extract_tables(doc) -> dict[str, str]:
    """
    Walk all tables in the document and extract Field → Value pairs.

    Handles:
    - 2-column tables: [Field, Value]
    - 3-column tables: [Field, Value, Status] — status column ignored
    - Header rows are skipped
    - Multi-line cell content is joined with newlines for multiline keys,
      or collapsed to a single line for simple values
    - Duplicate keys: last value wins (later sections override earlier ones)
    """
    data: dict[str, str] = {}
    warnings: list[str] = []

    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) < 2:
                continue

            raw_key = cells[0].text.strip()
            raw_val = _cell_text(cells[1])

            if not raw_key or not raw_val:
                continue

            # Skip header rows
            if _is_header_row([raw_key]):
                continue

            canonical = _resolve_key(raw_key)
            if not canonical:
                # Not a recognised field — skip silently
                # (the doc has many informational rows we don't need)
                continue

            # For multiline fields preserve full text; for others collapse.
            # Custom tool PL/SQL bodies, code, and instructions need newlines
            # preserved regardless of N, so check the dynamic pattern too.
            is_multiline = (
                canonical in _MULTILINE_KEYS
                or re.match(r"^custom_tool_\d+_(plsqlbody|code|instruction|"
                            r"pythonscript|dataquery|parameters|paramspec)$", canonical)
            )
            if is_multiline:
                value = raw_val  # keep newlines
            else:
                # Collapse to single line, strip status badges if present
                value = " ".join(raw_val.splitlines()).strip()
                # Strip trailing status badges like "  FOUND" "  MISSING"
                value = re.sub(r"\s+(FOUND|MISSING|PARTIAL|READY|NEW)\s*✓?\s*$",
                               "", value, flags=re.IGNORECASE).strip()

            if _CODE_FIELD_PAT.match(canonical):
                value = _normalize_code_punctuation(value)

            if value:
                data[canonical] = value

    return data, warnings


def _extract_sql_tables(data: dict[str, str]) -> list[str]:
    """
    Build the SQL tables list from whatever table references we found.

    Sources (in priority order):
    1. NL2SQL Object List / Underlying Table N fields — N can be any count,
       not just 1 and 2. A schema like ACME_CORP with 6 source tables works
       the same way as a 2-table document; the field count simply scales.
    2. sql_tables (comma-separated — used by "NL2SQL Object List" single field)
    3. sql_schema + any table names we can infer
    """
    tables = []

    # Direct numbered table fields (Underlying Table 1, 2, 3, ... N) —
    # collect ALL sql_table_<n> keys present, in numeric order, not a
    # hardcoded pair.
    numbered_keys = []
    for key in data:
        m = re.match(r"^sql_table_(\d+)$", key)
        if m:
            numbered_keys.append((int(m.group(1)), key))
    for _, key in sorted(numbered_keys):
        val = data.get(key, "").strip()
        if val:
            val = re.sub(r"\s+", " ", val).upper().strip()
            if val not in tables:
                tables.append(val)

    # Comma/semicolon-separated list — e.g. a single "NL2SQL Object List" field
    if not tables and data.get("sql_tables"):
        raw = data["sql_tables"]
        for part in re.split(r"[,;]+", raw):
            part = part.strip().upper()
            if part and part not in tables:
                tables.append(part)

    return tables


def _extract_custom_tools(data: dict[str, str]) -> list[dict]:
    """
    Assemble custom function-type tools from "Custom Tool N <field>" cells.

    Two ways to define a tool — pick whichever fits:

    1. PL/SQL Body (or Code) — paste an already-tested, complete
       CREATE OR REPLACE FUNCTION/PROCEDURE/PACKAGE block verbatim. Used
       byte-for-byte, no generation involved. Best for hand-tuned business
       logic that doesn't fit a generic shape.

    2. Python Script + Data Query — for OML4Py (Embedded Python Execution)
       tools specifically. Supply just the Python function body, the SQL
       query that builds its JSON data payload, and an optional Parameters
       spec; the builder generates the full wrapper itself, including the
       auto-refreshing OML auth token handling (token expires roughly
       every 60 minutes — the generated wrapper fetches a fresh one on
       every call so this is never a manual step) and the script
       registration. See core.sql_builder._build_oml_python_tool.

    Function Name and Instruction are optional in both paths — Function
    Name falls back to "<tool name>_FN", Instruction to a generic
    placeholder (a real instruction significantly improves routing
    accuracy, but its absence won't break anything).

    Returns a list of dicts. PL/SQL-body tools carry "raw_plsql"; OML4Py
    tools carry "python_script"/"data_query"/"pyqscript_name"/
    "credential_name"/"param_spec" — sql_builder dispatches on which keys
    are present.
    """
    field_pat = re.compile(
        r"^custom_tool_(\d+)_(name|functionname|instruction|plsqlbody|code|type|inputs"
        r"|pythonscript|dataquery|pyqscriptname|credentialname|parameters|paramspec)$"
    )
    groups: dict[str, dict] = {}
    for key, val in data.items():
        m = field_pat.match(key)
        if not m:
            continue
        n, field = m.group(1), m.group(2)
        groups.setdefault(n, {})[field] = val

    tools = []
    skipped_placeholders = []
    for n in sorted(groups, key=int):
        g = groups[n]
        name = (g.get("name") or "").strip().upper()
        if not name:
            continue

        function_name = (g.get("functionname") or f"{name}_FN").strip().upper()
        instruction = (g.get("instruction") or
                       f"Custom tool {name} — see implementation for details.").strip()
        tool_type = (g.get("type") or "CUSTOM").strip().upper()

        body = (g.get("plsqlbody") or g.get("code") or "").strip()
        python_script = (g.get("pythonscript") or "").strip()
        data_query = (g.get("dataquery") or "").strip()

        is_placeholder = bool(body) and (
            re.match(r"^\s*(PASTE|TODO|TBD|PLACEHOLDER|REPLACE\s+WITH)\b", body, re.IGNORECASE)
            or not re.search(r"\bCREATE\s+(OR\s+REPLACE\s+)?(FUNCTION|PROCEDURE|PACKAGE)\b",
                              body, re.IGNORECASE)
        )

        if body and not is_placeholder:
            # Path 1: complete, ready-to-run PL/SQL supplied directly.
            tools.append({
                "name": name, "type": tool_type,
                "function_name": function_name, "instruction": instruction,
                "inputs": [], "raw_plsql": body,
            })
        elif python_script or data_query:
            # Path 2: builder generates the wrapper from these pieces.
            tools.append({
                "name": name, "type": tool_type or "PYTHON",
                "function_name": function_name, "instruction": instruction,
                "inputs": [],
                "python_script": python_script,
                "data_query": data_query,
                "pyqscript_name": (g.get("pyqscriptname") or "").strip(),
                "credential_name": (g.get("credentialname") or "").strip(),
                "param_spec": (g.get("parameters") or g.get("paramspec") or "").strip(),
            })
        else:
            # Neither a usable PL/SQL body nor enough to generate one —
            # placeholder field, or genuinely incomplete entry.
            skipped_placeholders.append(name)
            continue

    if skipped_placeholders:
        print(f"  ⚠  Skipped {len(skipped_placeholders)} custom tool(s) — incomplete or "
              f"placeholder definition (need either a real PL/SQL Body, or both a Python "
              f"Script and Data Query): " + ", ".join(skipped_placeholders))
    return tools


def _parse_comment_sql(sql_text: str) -> list[dict]:
    """
    Parse raw COMMENT ON TABLE/COLUMN SQL text into a list of structured dicts.

    Handles:
      COMMENT ON TABLE owner.table IS 'text';
      COMMENT ON COLUMN owner.table.column IS 'text';

    Oracle-style escaped single quotes ('') are unescaped to (') in the result.
    Returns list of dicts with keys: type, owner, table, column (optional), text.
    """
    results = []
    # Match COMMENT ON TABLE schema.table IS '...'  (text may span lines)
    table_pat = re.compile(
        r"COMMENT\s+ON\s+TABLE\s+(\w+)\.(\w+)\s+IS\s+'((?:[^']|'')*)'",
        re.IGNORECASE | re.DOTALL
    )
    # Match COMMENT ON COLUMN schema.table.column IS '...'
    col_pat = re.compile(
        r"COMMENT\s+ON\s+COLUMN\s+(\w+)\.(\w+)\.(\w+)\s+IS\s+'((?:[^']|'')*)'",
        re.IGNORECASE | re.DOTALL
    )
    for m in table_pat.finditer(sql_text):
        owner, table, text = m.group(1), m.group(2), m.group(3)
        results.append({
            "type": "table", "owner": owner.upper(),
            "table": table.upper(),
            "text": text.replace("''", "'").strip()
        })
    for m in col_pat.finditer(sql_text):
        owner, table, col, text = m.group(1), m.group(2), m.group(3), m.group(4)
        results.append({
            "type": "column", "owner": owner.upper(),
            "table": table.upper(), "column": col.upper(),
            "text": text.replace("''", "'").strip()
        })
    return results


def _populate_comments_from_sql(project: dict, sql_blocks: list[str]) -> int:
    """
    Parse COMMENT ON SQL blocks and populate facts.sql.comments.objects
    in the same structure that modules/comments.py uses natively.

    Returns the total number of comments imported.
    """
    facts = project.setdefault("facts", {})
    sql_facts = facts.setdefault("sql", {})
    comments = sql_facts.setdefault("comments", {
        "mode": "imported", "status": "approved",
        "objects": {}, "coverage": {}
    })
    objects = comments.setdefault("objects", {})
    count = 0

    for block in sql_blocks:
        if not block:
            continue
        for item in _parse_comment_sql(block):
            if not item.get("text"):
                continue
            key = f"{item['owner']}.{item['table']}"
            obj = objects.setdefault(key, {"table_comment": "", "columns": {}})
            if item["type"] == "table":
                obj["table_comment"] = item["text"]
                count += 1
            elif item["type"] == "column":
                obj.setdefault("columns", {})[item["column"]] = item["text"]
                count += 1

    if count:
        comments["mode"]   = "imported"
        comments["status"] = "approved"
        # Inline coverage calculation (mirrors comments.calculate_coverage)
        table_total  = len(objects)
        table_with   = sum(1 for o in objects.values() if o.get("table_comment"))
        col_total    = sum(len(o.get("columns", {})) for o in objects.values())
        col_with     = sum(
            sum(1 for v in o.get("columns", {}).values() if v)
            for o in objects.values()
        )
        comments["coverage"] = {
            "table_total": table_total, "table_with_comments": table_with,
            "col_total": col_total,   "col_with_comments": col_with,
        }

    return count


# ─────────────────────────────────────────────────────────────────────────────
# Parse docx → data dict
# ─────────────────────────────────────────────────────────────────────────────

def parse_docx(path: str | Path) -> tuple[dict[str, str], list[str]]:
    """
    Parse a Word document (.docx) into the canonical data dict format.

    Returns:
        (data, warnings)
        data     — {canonical_key: value_string}
        warnings — list of human-readable issues found during parsing
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx is required for Word document import.\n"
            "Install it with:  pip install python-docx"
        )

    path = Path(path).expanduser().resolve()
    doc = Document(str(path))

    data, warnings = _extract_tables(doc)

    # Post-process: build sql tables list from component fields
    tables = _extract_sql_tables(data)
    if tables:
        data["sql_tables"] = ", ".join(tables)
        # Clean up component keys — build_project doesn't need them.
        # Dynamic: removes sql_table_1 through sql_table_N, whatever N is,
        # not just a hardcoded 1/2 pair.
        for key in [k for k in data if re.match(r"^sql_table_\d+$", k)]:
            data.pop(key, None)

    # Post-process: if sql_schema is present but sql_tables already has
    # schema-qualified names, sql_schema is just informational — keep it
    # but don't overwrite the table list.

    # Post-process: rag_url — if "N/A" or similar, clear it
    rag_url = data.get("rag_url", "")
    if rag_url.lower() in ("n/a", "na", "none", "tbd", "-", ""):
        data.pop("rag_url", None)
    elif not rag_url.startswith("http"):
        # May have been set from a non-URL field — clear it
        data.pop("rag_url", None)
        warnings.append(f"RAG URL '{rag_url}' does not look like a URL — cleared")

    # Post-process: data_source — infer from what's present if not explicit
    if "data_source" not in data:
        has_rag = bool(data.get("rag_url") or data.get("rag_subject"))
        has_sql = bool(data.get("sql_tables") or data.get("sql_table_1"))
        if has_rag and has_sql:
            data["data_source"] = "both"
        elif has_rag:
            data["data_source"] = "documents"
        elif has_sql:
            data["data_source"] = "tables"

    return data, warnings


# ─────────────────────────────────────────────────────────────────────────────
# Extend build_project for docx-specific fields
# ─────────────────────────────────────────────────────────────────────────────

def build_project_from_docx(data: dict[str, str], cfg, schema_default: str = "") -> dict:
    """
    Build a project dict from parsed Word doc data.

    Delegates to project_import.build_project() for all standard fields,
    then adds docx-specific fields (tool descriptions, agent description,
    chat/embed model, sql schema).
    """
    project = build_project(data, cfg, schema_default)
    facts   = project["facts"]
    names   = facts.setdefault("names", {})

    # SQL tool description (stored in names for codegen access)
    if data.get("sql_tool_description"):
        names["sql_tool_description"] = data["sql_tool_description"].strip()

    # RAG tool description
    if data.get("rag_tool_description"):
        names["rag_tool_description"] = data["rag_tool_description"].strip()

    # Agent description
    if data.get("agent_description"):
        project["agent_description"] = data["agent_description"].strip()

    # LLM model preferences (override config defaults if set in doc)
    # Store in facts.llm so spec_builder can prefer doc values over config.ini
    llm = project.setdefault("facts", {}).setdefault("llm", {})
    if data.get("chat_model"):
        llm["chat_model"]   = data["chat_model"].strip()
        names["chat_model"] = llm["chat_model"]
    if data.get("embed_model"):
        llm["embed_model"]   = data["embed_model"].strip()
        names["embed_model"] = llm["embed_model"]
    if data.get("llm_temperature"):
        try:
            llm["temperature"] = float(data["llm_temperature"].strip())
        except ValueError:
            pass
    if data.get("llm_max_tokens"):
        try:
            llm["max_tokens"] = int(str(data["llm_max_tokens"]).strip())
        except ValueError:
            pass

    # SQL target schema (may differ from the owning schema)
    if data.get("sql_schema"):
        names["sql_schema"] = data["sql_schema"].strip().upper()

    # RAG bucket (informational — URL is the authoritative source)
    if data.get("rag_bucket"):
        names["rag_bucket"] = data["rag_bucket"].strip()

    # NL2SQL comments — parse COMMENT ON SQL blocks from the Word doc.
    # The doc may split them across two cells (table-level / column-level)
    # or combine them into one. Collect all blocks and parse together.
    sql_blocks = []
    for key in ("nl2sql_comments_sql", "nl2sql_comments_sql_col", "nl2sql_comments"):
        val = data.get(key, "").strip()
        if val:
            sql_blocks.append(val)
    if sql_blocks:
        n = _populate_comments_from_sql(project, sql_blocks)
        if n:
            import sys as _sys
            print(f"  →  Imported {n} NL2SQL comment(s) from document")

    # Custom function-type tools (PL/SQL or OML4Py-wrapped) — "Custom Tool N
    # <field>" cells. These carry their own already-tested implementation
    # body, imported verbatim rather than synthesised from a template.
    custom_tools = _extract_custom_tools(data)
    if custom_tools:
        existing = facts.setdefault("analysis_tools", [])
        existing_names = {t.get("name") for t in existing}
        added = 0
        for ct in custom_tools:
            if ct["name"] not in existing_names:
                existing.append(ct)
                added += 1
        if added:
            print(f"  →  Imported {added} custom tool(s) from document: "
                  + ", ".join(t["name"] for t in custom_tools))

    if data.get("error_log_table"):
        facts["error_log_table"] = data["error_log_table"].strip().upper()

    # Mark as imported from docx for run log
    project["_imported_from_docx"] = True

    return project


# ─────────────────────────────────────────────────────────────────────────────
# Summary display — extends the CSV summary with docx-specific fields
# ─────────────────────────────────────────────────────────────────────────────

def print_docx_summary(project: dict, data: dict, warnings: list[str], display) -> None:
    """Print a full import summary including docx-specific fields."""
    C = display.C

    # Base summary (shared with CSV import)
    print_import_summary(project, [], display)

    # Additional docx-specific fields
    facts = project["facts"]
    names = facts.get("names", {})

    print(f"  {C.BOLD}Additional fields from document{C.RESET}")
    print(f"  {'─'*60}")

    def row(label, value):
        val_str = str(value)[:100] if value else f"{C.DIM}(not set){C.RESET}"
        print(f"  {label:<30}: {val_str}")

    row("Agent description",    project.get("agent_description"))
    row("SQL tool description", (names.get("sql_tool_description") or "")[:80])
    row("RAG tool description", (names.get("rag_tool_description") or "")[:80])
    row("Chat model",           names.get("chat_model"))
    row("Embed model",          names.get("embed_model"))
    llm_facts = project.get("facts", {}).get("llm", {})
    if llm_facts.get("temperature"):
        row("Temperature",      str(llm_facts["temperature"]))
    if llm_facts.get("max_tokens"):
        row("Max tokens",       str(llm_facts["max_tokens"]))
    row("SQL target schema",    names.get("sql_schema"))
    row("RAG bucket",           names.get("rag_bucket"))

    # Show instruction previews
    agent_role = facts.get("agent_role", "")
    if agent_role:
        preview = agent_role.replace("\n", " ")[:100]
        print(f"  {'Agent instructions':<30}: {preview}{'...' if len(agent_role) > 100 else ''}")

    task_instr = facts.get("task", {}).get("instruction", "")
    if task_instr:
        preview = task_instr.replace("\n", " ")[:100]
        print(f"  {'Task instruction':<30}: {preview}{'...' if len(task_instr) > 100 else ''}")

    display.blank()
    if warnings:
        print(f"  {C.YELLOW}Warnings:{C.RESET}")
        for w in warnings:
            print(f"    {C.YELLOW}⚠{C.RESET}  {w}")
        display.blank()

    print(f"  {'─'*60}")


# ─────────────────────────────────────────────────────────────────────────────
# Main entry point
# ─────────────────────────────────────────────────────────────────────────────

def run_import_docx(cfg, clients: dict, display) -> Optional[dict]:
    """
    Prompt for a Word doc path, parse it, show summary, and return the
    project dict ready for _discovery_loop (starting at Step 7) — or
    None if the user cancelled or validation failed.
    """
    C = display.C
    display.head("IMPORT PROJECT FROM WORD DOCUMENT")
    display.blank()
    print(f"  {C.DIM}Reads the standard Select AI Agent configuration document (.docx).{C.RESET}")
    print(f"  {C.DIM}All two-column tables (Field | Value) in the document are parsed.{C.RESET}")
    print(f"  {C.DIM}Multi-line fields (Agent Instructions, Task Instruction) are preserved.{C.RESET}")
    display.blank()

    try:
        raw_path = input("  Word document path (or q to cancel): ").strip()
    except (EOFError, KeyboardInterrupt):
        display.warn("Cancelled")
        return None

    if not raw_path or raw_path.lower() in ("q", "quit", "exit", "b", "back"):
        display.warn("Cancelled")
        return None

    path = Path(raw_path).expanduser().resolve()
    if not path.exists():
        display.err(f"File not found: {path}")
        return None
    if path.suffix.lower() not in (".docx",):
        display.err(f"Expected a .docx file, got '{path.suffix}'")
        display.info("For CSV files use Import from CSV (option 2)")
        return None

    # Parse
    display.info("Parsing Word document...")
    try:
        data, parse_warnings = parse_docx(path)
    except ImportError as ex:
        display.err(str(ex))
        return None
    except Exception as ex:
        display.err(f"Failed to read document: {ex}")
        return None

    if not data:
        display.err("No recognised fields found in the document.")
        display.info("Check the document has two-column tables with Field | Value columns.")
        return None

    display.ok(f"Parsed {len(data)} field(s) from document")

    # Build project
    # NOTE: cfg_module.get's fallback only fires when a key is entirely
    # ABSENT from the config file — configparser does not treat a key that
    # is present but explicitly blank (target_schema = , as set up for
    # direct-login mode — see [database] comments) the same as a missing
    # key. Without this explicit check, direct-login users would always
    # fall through to the literal "YOUR_SCHEMA" placeholder instead of
    # their own db_user. Same fix pattern already applied in
    # modules/preflight_schema.py for the identical underlying issue.
    raw_target_schema = cfg_module.get(cfg, "database", "target_schema", fallback="")
    schema_default = raw_target_schema or cfg_module.get(
        cfg, "database", "db_user", fallback="YOUR_SCHEMA"
    )
    try:
        project = build_project_from_docx(data, cfg, schema_default)
    except Exception as ex:
        display.err(f"Failed to build project from document data: {ex}")
        return None

    # Validate
    errors, val_warnings = validate_project(project)
    all_warnings = parse_warnings + val_warnings

    # Show summary
    print_docx_summary(project, data, all_warnings, display)

    if errors:
        display.err("Import cannot continue — required fields are missing:")
        for e in errors:
            display.err(f"  • {e}")
        display.info("Fill in the missing fields in the document and try again.")
        return None

    # Confirm
    display.blank()
    print(f"  {C.BOLD}Options:{C.RESET}")
    print(f"   1. Proceed    — go to Step 7 to confirm names and generate SQL")
    print(f"   2. Edit first — start the guided conversation with pre-loaded values")
    print(f"   3. Cancel")
    display.blank()
    try:
        choice = input("  Choice [1/2/3]: ").strip()
    except (EOFError, KeyboardInterrupt):
        display.warn("Cancelled")
        return None

    if choice == "3" or choice.lower() in ("q", "quit", "cancel"):
        display.warn("Import cancelled")
        return None

    if choice == "2":
        project["workflow"]["current_step"] = 2
        project["workflow"]["last_completed_step"] = 1
        display.info("Starting guided conversation. Your document values are pre-loaded as context.")

    # Check for existing project
    existing = state_module.projects_dir(cfg) / project["project_name"] / f"{project['project_name']}.json"
    if not existing.exists():
        # Legacy flat path check
        existing = state_module.projects_dir(cfg) / f"{project['project_name']}.json"

    if existing.exists():
        display.blank()
        display.warn(f"A project named '{project['project_name']}' already exists")
        try:
            overwrite = input("  Overwrite it? [y/N]: ").strip().lower()
        except (EOFError, KeyboardInterrupt):
            overwrite = "n"
        if overwrite not in ("y", "yes"):
            display.warn("Import cancelled — rename the project in the document and try again")
            return None

    # Save
    display.blank()
    state_module.save_project(cfg, project)
    run_log = state_module.open_run_log(cfg, project)
    run_log.log_section("PROJECT IMPORTED FROM WORD DOCUMENT")
    run_log.log(f"Source  : {path}")
    run_log.log(f"Project : {project['display_name']}  Schema: {project['schema']}")
    run_log.log(f"Fields  : {len(data)} parsed from document")

    display.ok(f"Project saved: {project['project_name']}")
    display.info(f"Run log: {run_log.run_path}")

    project["_run_log"] = run_log
    return project
